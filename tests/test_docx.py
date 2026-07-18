"""Phase 4b — DOCX format layering: python-docx extractor, format dispatch, and needles-
through-conversion (§7.3, §15.4). DOCX is read + written with python-docx (the [docx] extra /
[dev])."""

from __future__ import annotations

import io
from pathlib import Path

from docusearch.convert import convert_corpus, html_to_docx_bytes
from docusearch.ingest import extract_document, extract_docx


def _add_hyperlink(paragraph, url: str, text: str) -> None:  # type: ignore[no-untyped-def]
    # python-docx has no high-level hyperlink API; build the w:hyperlink relationship by hand.
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _make_docx(**kw) -> bytes:  # type: ignore[no-untyped-def]
    from docx import Document

    d = Document()
    d.core_properties.title = kw.get("title", "")
    d.add_heading("SPI Protocol Overview", level=1)
    d.add_paragraph("The nonce ZQX-7734-FRB configures the peripheral bus.")
    d.add_heading("Timing", level=2)
    d.add_paragraph("Setup and hold windows around the strobe.")
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Marker"
    table.rows[0].cells[1].text = "Note"
    table.rows[1].cells[0].text = "MLP-4242-XY"
    table.rows[1].cells[1].text = "per site"
    p = d.add_paragraph("See ")
    _add_hyperlink(p, "https://example.com/spec", "the spec")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_extract_docx_headings_paras_tables_links() -> None:
    doc = extract_docx(_make_docx(title="SPI Doc"))
    text = "\n".join(s.text for s in doc.segments)
    assert "ZQX-7734-FRB" in text  # prose under a heading
    assert "Setup and hold windows" in text
    # heading paths carry the section structure
    hpaths = {s.heading_path for s in doc.segments}
    assert any("SPI Protocol Overview" in h for h in hpaths)
    assert any("Timing" in h for h in hpaths)
    # table linearized, its needle recoverable, kind=table
    tbl = [s for s in doc.segments if s.kind == "table"]
    assert tbl and "MLP-4242-XY" in tbl[0].text and "|" in tbl[0].text
    # hyperlink captured
    assert any(lk.target == "https://example.com/spec" and lk.link_type == "docx_hyperlink"
               for lk in doc.links)
    assert doc.title == "SPI Doc"


def test_extract_document_dispatch_docx(tmp_path: Path) -> None:
    p = tmp_path / "a.docx"
    p.write_bytes(_make_docx())
    segs = extract_document(p, "docx").segments
    assert "ZQX-7734-FRB" in " ".join(s.text for s in segs)


def test_html_to_docx_roundtrip_preserves_needles_all_placements() -> None:
    # prose / code / table / image-alt nonces must all survive HTML -> DOCX -> extract (§15.4)
    html = (
        "<body><h1>SPI</h1>"
        "<p>Prose needle PRO-1000-XX configures the bus.</p>"
        "<pre><code>CONST = 'CODE-2000-YY'  # timing</code></pre>"
        "<table><tr><td>TAB-3000-ZZ</td><td>per site</td></tr></table>"
        '<img src="missing.png" alt="diagram IMG-4000-WW for the gyroscope">'
        "</body>"
    )
    doc = extract_docx(html_to_docx_bytes(html))
    text = "\n".join(s.text for s in doc.segments)
    for nonce in ("PRO-1000-XX", "CODE-2000-YY", "TAB-3000-ZZ", "IMG-4000-WW"):
        assert nonce in text, f"{nonce} lost in HTML->DOCX->extract"


def test_extract_docx_nested_table_not_dropped() -> None:
    # red-team H1: a table nested inside another table's cell must NOT be silently lost.
    from docx import Document

    d = Document()
    outer = d.add_table(rows=1, cols=1)
    cell = outer.rows[0].cells[0]
    cell.paragraphs[0].text = "outer OUT-1111-AA"
    inner = cell.add_table(rows=1, cols=2)
    inner.rows[0].cells[0].text = "NEST-2222-BB"
    inner.rows[0].cells[1].text = "inner value"
    buf = io.BytesIO()
    d.save(buf)
    text = "\n".join(s.text for s in extract_docx(buf.getvalue()).segments)
    assert "OUT-1111-AA" in text  # outer cell text
    assert "NEST-2222-BB" in text  # nested-table content survives (H1)


def test_extract_docx_hyperlink_in_table_cell() -> None:
    # red-team M1: a hyperlink whose only placement is inside a table cell must reach the graph.
    from docx import Document

    d = Document()
    t = d.add_table(rows=1, cols=1)
    _add_hyperlink(t.rows[0].cells[0].paragraphs[0], "https://example.com/incell", "in cell")
    buf = io.BytesIO()
    d.save(buf)
    links = extract_docx(buf.getvalue()).links
    assert any(lk.target == "https://example.com/incell" for lk in links)


def test_html_to_docx_emits_real_table_rows_distinct() -> None:
    # red-team M2: html_to_docx must emit a genuine DOCX table (not a flattened paragraph), so
    # row boundaries survive and extract_docx's table path is exercised.
    html = (
        "<body><h1>T</h1>"
        "<table><tr><td>R1C1-AAA</td><td>R1C2-BBB</td></tr>"
        "<tr><td>R2C1-CCC</td><td>R2C2-DDD</td></tr></table></body>"
    )
    doc = extract_docx(html_to_docx_bytes(html))
    tbl = [s for s in doc.segments if s.kind == "table"]
    assert tbl, "table segment must round-trip as kind=table, not a paragraph"
    # each source row is its own linearized row (not merged across the row boundary)
    rows = tbl[0].text.splitlines()
    assert any("R1C1-AAA | R1C2-BBB" in r for r in rows)
    assert any("R2C1-CCC | R2C2-DDD" in r for r in rows)
    assert not any("BBB R2C1" in r for r in rows)  # rows must not have flattened together


def test_convert_corpus_docx(tmp_path: Path) -> None:
    src = tmp_path / "html"
    (src / "sub").mkdir(parents=True)
    (src / "a.html").write_text("<body><h1>A</h1><p>needle ZQX-7734-FRB here.</p></body>", "utf-8")
    (src / "sub" / "b.html").write_text("<body><h1>B</h1><p>token MLP-4242-XY here.</p></body>", "utf-8")
    dst = tmp_path / "docx"
    result = convert_corpus(src, dst, fmt="docx")
    assert result.converted == 2 and not result.errors
    assert (dst / "a.docx").is_file() and (dst / "sub" / "b.docx").is_file()
    assert "ZQX-7734-FRB" in " ".join(s.text for s in extract_docx((dst / "a.docx").read_bytes()).segments)
