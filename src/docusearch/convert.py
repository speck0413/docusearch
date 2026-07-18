"""Derived-corpus builder for format-layering validation (Phase 4, §17).

Converts a baseline HTML corpus to another format so the golden + needle suites can prove that
retrieval survives the conversion — i.e. that the format's extractor (Phase 4a: PyMuPDF)
recovers the content that HTML ingest indexed. This is **harness-only**: ``reportlab`` is lazy-
imported and lives in the ``[dev]`` extra, never the runtime deps (ingesting real PDFs needs
only ``[pdf]``/PyMuPDF).

The conversion preserves the document's *tokens* (every needle, identifier, and word), not its
visual formatting — a PDF's text layer carries no code/table structure anyway, so the point of
the derived corpus is to check that content, not styling, makes the round trip.

Public surface:
    html_to_pdf_bytes(html) -> bytes                 # one HTML doc's text -> a PDF
    convert_corpus(src_dir, dst_dir, fmt="pdf") -> ConvertResult
"""

from __future__ import annotations

import contextlib
import io
import re
from collections.abc import Callable, Sequence
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING, cast
from xml.sax.saxutils import escape

from .ingest import ImageRef, Segment, extract_html

if TYPE_CHECKING:
    from .config import SourceConfig

_SUPPORTED = ("pdf", "docx", "md")
# Max embedded-image box on a US-letter page (612x792 pt) inside 1" margins, with headroom so a
# tall diagram (e.g. 450x990) is scaled down to fit instead of raising reportlab's LayoutError.
_PDF_IMG_MAX_W = 450.0
_PDF_IMG_MAX_H = 610.0


def _resolve_image(src: str, base_path: Path | str | None) -> Path | None:
    """Resolve an ``<img src>`` to a readable local image file, or None (external/missing/needle)."""
    if base_path is None:
        return None
    from .ingest import _resolve_local

    local = _resolve_local(src, Path(base_path))
    return local if local is not None and local.is_file() else None


def _doc_blocks(doc: object) -> list[tuple[str, object]]:
    """Interleave segments and images so each image is emitted **under its own heading path**, not
    dumped in a trailing loop under the last heading (which collapsed every image's locator, §7.3
    R-ING-6). Returns ``[("seg", Segment) | ("img", ImageRef)]`` in document order."""
    from collections import defaultdict

    by_head: dict[str, list[object]] = defaultdict(list)
    for im in doc.images:  # type: ignore[attr-defined]
        by_head[im.heading_path].append(im)
    out: list[tuple[str, object]] = []
    seen: set[int] = set()
    last: str | None = None
    for seg in doc.segments:  # type: ignore[attr-defined]
        if last is not None and seg.heading_path != last:
            out += [("img", im) for im in by_head.get(last, []) if id(im) not in seen]
            seen.update(id(im) for im in by_head.get(last, []))
        last = seg.heading_path
        out.append(("seg", seg))
    for ims in by_head.values():  # the last heading's images + any orphans
        out += [("img", im) for im in ims if id(im) not in seen]
        seen.update(id(im) for im in ims)
    return out


_MD_ESCAPE = re.compile(r"([\\`*_\[\]<>|~])")


def _md_escape(text: str) -> str:
    """Backslash-escape CommonMark metacharacters so identifiers survive as literal text — e.g.
    ``__init__`` would otherwise be parsed as emphasis and its underscores stripped. markdown-it
    unescapes them back at ingest, so the tokens are recovered intact."""
    return _MD_ESCAPE.sub(r"\\\1", text)


def html_to_pdf_bytes(
    html: str,
    *,
    content_selector: str = "",
    strip_selectors: Sequence[str] = (),
    base_path: Path | str | None = None,
) -> bytes:
    """Render one HTML document's extracted text to a single-column PDF (reportlab).

    Every segment becomes a wrapped paragraph (heading paths as sub-headings), so all text
    tokens survive; PyMuPDF then recovers them at ingest. XML-special characters are escaped so
    reportlab's paragraph markup can't drop code containing ``<``/``>``/``&``.

    ``content_selector``/``strip_selectors`` mirror the source's ingest config so the derived
    PDF carries the same **cleaned** article text the HTML store indexes — not framework chrome.
    ``base_path`` (the source HTML file) lets real ``<img>`` files be **embedded** so they survive
    the round trip and can be vision-enriched (R-ING-6); a missing/external image falls back to its
    alt text.
    """
    from reportlab.lib.pagesizes import letter  # lazy: [dev] harness dependency
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import Image as RLImage
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    doc = extract_html(html, content_selector=content_selector, strip_selectors=list(strip_selectors))
    styles = getSampleStyleSheet()
    story: list[object] = []
    if doc.title:
        story.append(Paragraph(escape(doc.title), styles["Title"]))
    last_heading: str | None = None
    for kind, item in _doc_blocks(doc):
        if item.heading_path and item.heading_path != last_heading:  # type: ignore[attr-defined]
            story.append(Paragraph(escape(item.heading_path), styles["Heading2"]))  # type: ignore[attr-defined]
            last_heading = item.heading_path  # type: ignore[attr-defined]
        if kind == "seg":
            seg = cast(Segment, item)
            # newlines -> spaces so code lines don't merge into an unbroken token run
            story.append(Paragraph(escape(seg.text.replace("\n", " ")), styles["BodyText"]))
            story.append(Spacer(1, 6))
            continue
        img = cast(ImageRef, item)
        local = _resolve_image(img.src, base_path)
        if local is not None:  # embed the real image so it survives + can be vision-enriched
            try:
                iw, ih = ImageReader(str(local)).getSize()
                if iw > 0 and ih > 0:
                    # fit within the page box on BOTH axes (never upscale) so a tall diagram
                    # doesn't overflow the frame and abort the whole PDF build (LayoutError).
                    scale = min(_PDF_IMG_MAX_W / iw, _PDF_IMG_MAX_H / ih, 1.0)
                    story.append(RLImage(str(local), width=iw * scale, height=ih * scale))
                    story.append(Spacer(1, 6))
            except Exception:  # noqa: BLE001 - a bad image falls back to its alt text below
                pass
        # Keep alt/caption text too (searchable; also the needle channel — §15.2 hides nonces here).
        caption = " ".join(t for t in (img.alt, img.caption) if t).strip()
        if caption:
            story.append(Paragraph(escape(caption), styles["Italic"]))
            story.append(Spacer(1, 6))
    if not story:  # never emit an empty PDF — keep the doc present for the audit counts
        story.append(Paragraph(escape(doc.title or "(no extractable text)"), styles["BodyText"]))
    buf = io.BytesIO()
    SimpleDocTemplate(buf, pagesize=letter, title=doc.title or "").build(story)
    return buf.getvalue()


def _add_docx_table(out: object, linearized: str) -> None:
    """Rebuild a real DOCX table from a ``a | b`` / newline-per-row linearized segment, so rows
    stay distinct (a flattened paragraph merges them) and ``extract_docx``'s table path is
    exercised end-to-end by the derived corpus (§15.4)."""
    grid = [row.split(" | ") for row in linearized.split("\n") if row.strip()]
    if not grid:
        return
    ncols = max(len(cells) for cells in grid)
    table = out.add_table(rows=len(grid), cols=ncols)  # type: ignore[attr-defined]
    for ri, cells in enumerate(grid):
        for ci, value in enumerate(cells):
            table.rows[ri].cells[ci].text = value


def html_to_docx_bytes(
    html: str,
    *,
    content_selector: str = "",
    strip_selectors: Sequence[str] = (),
    base_path: Path | str | None = None,
) -> bytes:
    """Render one HTML document's extracted text to a DOCX (python-docx).

    Mirrors :func:`html_to_pdf_bytes`: heading paths become ``Heading`` paragraphs, segments become
    paragraphs (tables as real DOCX tables), and real ``<img>`` files are **embedded** (via
    ``base_path``) so they survive + can be vision-enriched (R-ING-6); alt/caption text is also
    emitted (searchable + the needle channel). ``content_selector``/``strip_selectors`` mirror the
    source's ingest config so the derived DOCX carries the same cleaned content.
    """
    from docx import Document  # lazy: python-docx ([docx] extra / [dev])
    from docx.shared import Inches

    doc = extract_html(html, content_selector=content_selector, strip_selectors=list(strip_selectors))
    out = Document()
    if doc.title:
        out.core_properties.title = doc.title
        out.add_heading(doc.title, level=1)
    last_heading: str | None = None
    for kind, item in _doc_blocks(doc):
        if item.heading_path and item.heading_path != last_heading:  # type: ignore[attr-defined]
            out.add_heading(item.heading_path, level=2)  # type: ignore[attr-defined]
            last_heading = item.heading_path  # type: ignore[attr-defined]
        if kind == "seg":
            seg = cast(Segment, item)
            if seg.kind == "table":
                _add_docx_table(out, seg.text)  # a real DOCX table, so row boundaries survive
            else:
                # newlines -> spaces so code lines don't merge into an unbroken token run
                out.add_paragraph(seg.text.replace("\n", " "))
            continue
        img = cast(ImageRef, item)
        local = _resolve_image(img.src, base_path)
        if local is not None:  # embed the real image so it survives + can be vision-enriched
            with contextlib.suppress(Exception):  # a bad image falls back to its alt text below
                out.add_picture(str(local), width=Inches(4))
        caption = " ".join(t for t in (img.alt, img.caption) if t).strip()
        if caption:
            out.add_paragraph(caption)
    if not doc.segments and not doc.images:  # never emit an empty DOCX (keep it in audit counts)
        out.add_paragraph(doc.title or "(no extractable text)")
    buf = io.BytesIO()
    out.save(buf)
    return buf.getvalue()


def _md_table(linearized: str) -> list[str]:
    """A ``a | b`` / newline-per-row segment -> GFM table lines (header + separator + rows) so
    ``extract_md`` re-parses it as a real table."""
    grid = [[_md_escape(c) for c in row.split(" | ")] for row in linearized.split("\n") if row.strip()]
    if not grid:
        return []
    ncols = max(len(r) for r in grid)
    out = ["| " + " | ".join((grid[0] + [""] * ncols)[:ncols]) + " |",
           "| " + " | ".join(["---"] * ncols) + " |"]
    out += ["| " + " | ".join((r + [""] * ncols)[:ncols]) + " |" for r in grid[1:]]
    return out


def html_to_md_bytes(
    html: str,
    *,
    content_selector: str = "",
    strip_selectors: Sequence[str] = (),
    base_path: Path | str | None = None,
) -> bytes:
    """Render one HTML document's extracted content to Markdown (CommonMark + GFM).

    Heading paths become ``##`` headings, body becomes paragraphs, code becomes fenced blocks,
    tables become GFM tables, and real ``<img>`` files are embedded as ``![alt](data:...)`` data
    URIs (self-contained, so the image survives + can be vision-enriched, R-ING-6); a missing image
    keeps its alt text as ``![alt]()`` so image-placement needles survive. ``markdown-it-py`` at
    ingest recovers all of it.
    """
    import base64

    doc = extract_html(html, content_selector=content_selector, strip_selectors=list(strip_selectors))
    lines: list[str] = []
    if doc.title:
        lines += [f"# {doc.title}", ""]
    last_heading: str | None = None
    for kind, item in _doc_blocks(doc):
        if item.heading_path and item.heading_path != last_heading:  # type: ignore[attr-defined]
            # escape the heading text too, so `_`/`*` in a heading aren't parsed as emphasis
            lines += [f"## {_md_escape(item.heading_path)}", ""]  # type: ignore[attr-defined]
            last_heading = item.heading_path  # type: ignore[attr-defined]
        if kind == "seg":
            seg = cast(Segment, item)
            if seg.kind == "code":
                lines += ["```", seg.text, "```", ""]  # fenced code is literal — never escaped
            elif seg.kind == "table":
                lines += [*_md_table(seg.text), ""]
            else:  # prose: escape CommonMark metacharacters so identifiers (__init__) survive
                lines += [_md_escape(seg.text.replace("\n", " ")), ""]
            continue
        img = cast(ImageRef, item)
        alt = _md_escape(" ".join(t for t in (img.alt, img.caption) if t).strip())
        local = _resolve_image(img.src, base_path)
        src = ""
        if local is not None:  # embed the real image inline as a data URI (self-contained)
            with contextlib.suppress(Exception):
                ext = local.suffix.lstrip(".").lower() or "png"
                media = "jpeg" if ext in ("jpg", "jpeg") else ext
                b64 = base64.b64encode(local.read_bytes()).decode("ascii")
                src = f"data:image/{media};base64,{b64}"
        lines.append(f"![{alt}]({src})")
        lines.append("")
    if not doc.segments and not doc.images:
        lines.append(doc.title or "(no extractable text)")
    return ("\n".join(lines) + "\n").encode("utf-8")


def _render_bytes(
    html: str,
    fmt: str,
    *,
    content_selector: str = "",
    strip_selectors: Sequence[str] = (),
    base_path: Path | str | None = None,
) -> bytes:
    """Dispatch to the format's HTML->bytes renderer (the derived-corpus analog of the ingest
    extractor dispatch). A new format adds one branch here plus its ``html_to_<fmt>_bytes``."""
    kw = {"content_selector": content_selector, "strip_selectors": strip_selectors, "base_path": base_path}
    if fmt == "pdf":
        return html_to_pdf_bytes(html, **kw)  # type: ignore[arg-type]
    if fmt == "docx":
        return html_to_docx_bytes(html, **kw)  # type: ignore[arg-type]
    if fmt == "md":
        return html_to_md_bytes(html, **kw)  # type: ignore[arg-type]
    raise ValueError(f"unsupported target format {fmt!r}; supported: {_SUPPORTED}")


@dataclass
class ConvertResult:
    converted: int = 0
    errors: list[tuple[str, str]] = field(default_factory=list)  # (path, message)


def convert_corpus(src_dir: Path | str, dst_dir: Path | str, *, fmt: str = "pdf") -> ConvertResult:
    """Convert every HTML file under ``src_dir`` to ``fmt`` at the mirrored relative path under
    ``dst_dir`` (e.g. ``a/b.html`` -> ``a/b.pdf``). Returns audit counts; a per-file failure is
    recorded and the run continues."""
    if fmt not in _SUPPORTED:
        raise ValueError(f"unsupported target format {fmt!r}; supported: {_SUPPORTED}")
    src, dst = Path(src_dir), Path(dst_dir)
    result = ConvertResult()
    for html_path in sorted(src.rglob("*.htm*")):
        rel = html_path.relative_to(src).with_suffix(f".{fmt}")
        out = dst / rel
        try:
            html = html_path.read_bytes().decode("utf-8", errors="replace")
            data = _render_bytes(html, fmt, base_path=html_path)
            out.parent.mkdir(parents=True, exist_ok=True)
            out.write_bytes(data)
            result.converted += 1
        except Exception as err:  # noqa: BLE001 - one bad file must not abort the batch
            result.errors.append((str(html_path), f"{type(err).__name__}: {err}"))
    return result


def convert_source(
    source: SourceConfig,
    dst_dir: Path | str,
    *,
    fmt: str = "pdf",
    progress: Callable[[int, int], None] | None = None,
) -> ConvertResult:
    """Convert exactly the files a source ingests — ``iter_files`` applies the same include/
    exclude globs (R-REUSE-2) — into ``fmt`` under ``dst_dir``, mirroring relative paths and
    applying the source's ``content_selector``/``strip_selectors`` so the derived corpus carries
    the same cleaned content the HTML store indexes. This is how a real ingestion is *altered
    for the format under test* (§15.4 / R-PROC-8). Per-file failures are recorded, not fatal.
    """
    if fmt not in _SUPPORTED:
        raise ValueError(f"unsupported target format {fmt!r}; supported: {_SUPPORTED}")
    from .ingest import iter_files  # lazy: keeps convert import light

    src_root, dst = Path(source.location), Path(dst_dir)
    files = list(iter_files(source.location, source.include, source.exclude))
    total = len(files)
    result = ConvertResult()
    for i, path in enumerate(files, 1):
        try:
            rel = path.relative_to(src_root).with_suffix(f".{fmt}")
            out = dst / rel
            html = path.read_bytes().decode("utf-8", errors="replace")
            data = _render_bytes(
                html,
                fmt,
                content_selector=source.content_selector,
                strip_selectors=source.strip_selectors,
                base_path=path,
            )
            out.parent.mkdir(parents=True, exist_ok=True)
            out.write_bytes(data)
            result.converted += 1
        except Exception as err:  # noqa: BLE001 - one bad file must not abort the batch
            result.errors.append((str(path), f"{type(err).__name__}: {err}"))
        if progress is not None:
            progress(i, total)
    return result
